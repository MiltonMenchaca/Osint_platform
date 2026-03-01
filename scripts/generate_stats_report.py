from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from datetime import datetime

def create_stats_report():
    # Crear documento
    doc = Document()
    
    # Estilos básicos
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Título Principal
    title = doc.add_heading('Actividad: Introducción a la Estadística y Muestreo', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Información del Alumno
    doc.add_paragraph('')
    p = doc.add_paragraph()
    p.add_run('Alumno: ').bold = True
    p.add_run('Milton Emilio Menchaca Manero\n')
    p.add_run('Universidad: ').bold = True
    p.add_run('Tecmilenio\n')
    p.add_run('Materia: ').bold = True
    p.add_run('Estadística / Ciberseguridad (OSINT)\n')
    p.add_run('Fecha: ').bold = True
    p.add_run('12 de Febrero')
    
    doc.add_paragraph('---')

    # Tema Seleccionado
    doc.add_heading('Tema Seleccionado', level=1)
    p = doc.add_paragraph()
    runner = p.add_run('"Percepción de Privacidad y Exposición de Huella Digital (OSINT)"')
    runner.bold = True
    runner.italic = True
    
    doc.add_paragraph(
        'Relacionado con el proyecto de clase (Plataforma OSINT), el objetivo fue medir qué tanta '
        'información personal exponen voluntariamente los usuarios en sus redes sociales y qué tan '
        'conscientes son de los riesgos de ingeniería social.'
    )
    
    # Desarrollo
    doc.add_heading('Desarrollo de la Actividad', level=1)
    
    # Pregunta 1
    doc.add_heading('1. ¿Qué tipo de muestreo aplicaste?', level=2)
    p = doc.add_paragraph()
    p.add_run('Se aplicó un ').text
    p.add_run('Muestreo por Conveniencia (No Probabilístico).').bold = True
    
    doc.add_paragraph(
        'Este método consiste en seleccionar una muestra de la población por el hecho de que sea accesible. '
        'En este caso, se seleccionaron individuos de fácil acceso para el investigador: compañeros de clase '
        'de la Universidad Tecmilenio, familiares directos y amigos cercanos contactados a través de redes '
        'sociales (WhatsApp y Telegram).'
    )
    
    # Pregunta 2
    doc.add_heading('2. ¿Cómo lo llevaste a cabo?', level=2)
    doc.add_paragraph('El proceso siguió los siguientes pasos:')
    
    # Lista numerada
    doc.add_paragraph('Diseño del Instrumento:', style='List Number').runs[0].bold = True
    p = doc.add_paragraph('Se elaboró un cuestionario breve en Google Forms con 5 preguntas clave:')
    p.paragraph_format.left_indent = Pt(20)
    
    bullets = [
        '¿Tienes tu perfil de Instagram/Facebook público o privado?',
        '¿Publicas tu ubicación en tiempo real en tus historias?',
        '¿Usas la misma contraseña para múltiples servicios?',
        '¿Has buscado alguna vez tu propio nombre en Google para ver qué aparece?',
        'Nivel de estudios (para segmentación básica).'
    ]
    for b in bullets:
        p = doc.add_paragraph(b, style='List Bullet')
        p.paragraph_format.left_indent = Pt(40)

    p = doc.add_paragraph('Distribución:', style='List Number')
    p.runs[0].bold = True
    p.add_run(' Se compartió el enlace del formulario a través de grupos de WhatsApp de la universidad y grupos familiares.')
    
    p = doc.add_paragraph('Recolección:', style='List Number')
    p.runs[0].bold = True
    p.add_run(' Se obtuvieron respuestas durante un periodo de 24 horas.')

    # Pregunta 3
    doc.add_heading('3. ¿La muestra fue representativa?', level=2)
    p = doc.add_paragraph()
    p.add_run('No, la muestra no es estadísticamente representativa de la población general').bold = True
    p.add_run(', y esto es esperado en un muestreo por conveniencia.')
    
    items = [
        ('Sesgo Geográfico y Social:', ' La mayoría de los encuestados pertenecen a un círculo social similar (estudiantes universitarios de Tecmilenio o residentes de la misma ciudad), lo que no refleja la realidad de otros sectores demográficos.'),
        ('Sesgo Tecnológico:', ' Al distribuirse por medios digitales, se excluye a personas sin acceso a internet o redes sociales.'),
        ('Tamaño de la Muestra:', ' El número de encuestados es pequeño para generalizar resultados a nivel nacional o global.')
    ]
    
    for title, content in items:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(title).bold = True
        p.add_run(content)

    doc.add_paragraph(
        'Sin embargo, para los fines exploratorios de la clase y validar hipótesis sobre el comportamiento '
        'de seguridad en nuestro entorno inmediato, los datos son válidos cualitativamente.'
    )

    # Pregunta 4
    doc.add_heading('4. ¿Qué posibles problemas puede tener su método?', level=2)
    doc.add_paragraph('El método de muestreo por conveniencia presenta las siguientes limitaciones en este contexto:')
    
    problems = [
        ('Sesgo de Autoselección:', ' Solo responden las personas que tienen interés o tiempo, lo que puede dejar fuera a quienes son más celosos de su privacidad (y que precisamente no contestarían una encuesta sobre privacidad).'),
        ('Respuestas Socialmente Deseables:', ' Al ser familiares o amigos, pueden mentir sobre sus hábitos de seguridad (ej. decir que usan contraseñas seguras cuando no es así) para "quedar bien" con el investigador.'),
        ('Falta de Generalización:', ' No se pueden hacer inferencias matemáticas precisas (como márgenes de error o intervalos de confianza) sobre la población total basándose en estos resultados.')
    ]

    for title, content in problems:
        p = doc.add_paragraph(style='List Number')
        p.add_run(title).bold = True
        p.add_run(content)

    doc.add_paragraph('---')
    
    # Conclusión
    p = doc.add_paragraph()
    p.add_run('Conclusión para el Proyecto OSINT:').bold = True
    doc.add_paragraph(
        'A pesar de los sesgos, la encuesta reveló que un alto porcentaje de usuarios cercanos '
        '(nuestro "círculo de confianza") tiene prácticas de seguridad deficientes, lo que justifica '
        'la necesidad de herramientas como la Plataforma OSINT desarrollada en clase para auditoría y concientización.'
    )

    # Guardar archivo
    output_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'docs')
    output_path = os.path.join(output_dir, 'Actividad_Estadistica_Muestreo.docx')
    
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        
    doc.save(output_path)
    print(f"Documento generado exitosamente en: {output_path}")

if __name__ == "__main__":
    create_stats_report()
